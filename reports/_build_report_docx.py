
import os
from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "reports"
SOURCE_MD = Path(os.environ.get("REPORT_SOURCE_MD", OUT_DIR / "Bao_cao_do_an_CopyPro_AI_Copywriter_source.md"))
DOCX_PATH = Path(os.environ.get("REPORT_DOCX_PATH", OUT_DIR / "Bao_cao_do_an_CopyPro_AI_Copywriter_chi_tiet.docx"))
TXT_PATH = Path(os.environ.get("REPORT_TXT_PATH", OUT_DIR / "Bao_cao_do_an_CopyPro_AI_Copywriter_chi_tiet.txt"))
PAGE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120

REPORT_TITLE = "BÁO CÁO ĐỒ ÁN: COPYPRO - WEBSITE AI COPYWRITER"
REPORT_SUBTITLE = "Phân tích, thiết kế, xây dựng và triển khai hệ thống tạo nội dung marketing bằng AI"
REPORT_META = [
    ("Môn học", "NT114"),
    ("Tên project", "NT114.Q21-Web-Copy-Writing"),
    ("Tên sản phẩm", "CopyPro - AI Copywriter"),
    ("Sinh viên thực hiện", "Bùi Lê Huy Phước"),
    ("MSSV", "23521228"),
    ("Ngày lập tài liệu", "20/06/2026"),
    ("Phạm vi khảo sát", "Frontend Next.js, Backend Express.js, MongoDB/Mongoose models, REST API, AI services, fine-tuning, plagiarism, billing và admin portal"),
]


def set_run_font(run, name="Calibri", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rpr.rFonts.set(qn("w:ascii"), name)
    rpr.rFonts.set(qn("w:hAnsi"), name)
    rpr.rFonts.set(qn("w:cs"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_paragraph_spacing(paragraph, before=0, after=6, line=1.10):
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.allow_autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[min(idx, len(widths) - 1)])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_field_run(paragraph, field_code):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field_code
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(text)
    run._r.append(fld_end)


def configure_document(doc):
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    for style_name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        st = styles[style_name]
        st.font.name = "Calibri"
        st._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        st._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        st.font.size = Pt(size)
        st.font.color.rgb = RGBColor.from_string(color)
        st.font.bold = True
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.line_spacing = 1.10
    for style_name in ["List Bullet", "List Number"]:
        st = styles[style_name]
        st.font.name = "Calibri"
        st._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        st._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        st.font.size = Pt(11)
        st.paragraph_format.space_after = Pt(8)
        st.paragraph_format.line_spacing = 1.167
    hp = section.header.paragraphs[0]
    hp.text = "CopyPro AI Copywriter | Báo cáo đồ án NT114"
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph_spacing(hp, after=0)
    set_run_font(hp.runs[0], size=9, color=RGBColor(90, 90, 90), italic=True)
    fp = section.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(fp, after=0)
    r = fp.add_run("Trang ")
    set_run_font(r, size=9, color=RGBColor(90, 90, 90))
    add_field_run(fp, "PAGE")
    r2 = fp.add_run(" / ")
    set_run_font(r2, size=9, color=RGBColor(90, 90, 90))
    add_field_run(fp, "NUMPAGES")


def add_cover(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=12, after=4)
    r = p.add_run("ĐẠI HỌC / KHOA / BỘ MÔN")
    set_run_font(r, size=12, color=RGBColor(90, 90, 90), bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, after=28)
    r = p.add_run("ĐỒ ÁN MÔN HỌC NT114")
    set_run_font(r, size=14, color=RGBColor(31, 77, 120), bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=36, after=8)
    r = p.add_run(REPORT_TITLE)
    set_run_font(r, size=24, color=RGBColor(11, 37, 69), bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, after=28)
    r = p.add_run(REPORT_SUBTITLE)
    set_run_font(r, size=13, color=RGBColor(70, 70, 70), italic=True)
    table = doc.add_table(rows=len(REPORT_META), cols=2)
    table.style = "Table Grid"
    set_table_geometry(table, [2500, PAGE_WIDTH_DXA - 2500])
    for i, (label, value) in enumerate(REPORT_META):
        cells = table.rows[i].cells
        cells[0].text = label
        cells[1].text = value
        set_cell_shading(cells[0], "F2F4F7")
        for cell in cells:
            for paragraph in cell.paragraphs:
                set_paragraph_spacing(paragraph, after=2, line=1.10)
                for run in paragraph.runs:
                    set_run_font(run, size=10.5, bold=(cell is cells[0]))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=24, after=0)
    r = p.add_run("Tài liệu được tổng hợp từ mã nguồn project hiện tại trong workspace.")
    set_run_font(r, size=10.5, color=RGBColor(90, 90, 90), italic=True)
    doc.add_page_break()


def add_static_toc(doc):
    doc.add_heading("MỤC LỤC TÓM TẮT", level=1)
    items = [
        "TÓM TẮT ĐỒ ÁN",
        "Chương 1. GIỚI THIỆU ĐỀ TÀI",
        "Chương 2. CƠ SỞ LÝ THUYẾT VÀ CÔNG NGHỆ SỬ DỤNG",
        "Chương 3. PHÂN TÍCH THIẾT KẾ HỆ THỐNG",
        "Chương 4. XÂY DỰNG VÀ TRIỂN KHAI ỨNG DỤNG",
        "Chương 5. TỔNG KẾT",
        "TÀI LIỆU THAM KHẢO",
        "PHỤ LỤC A. MAPPING ROUTE THEO PROJECT HIỆN TẠI",
        "PHỤ LỤC B. DANH SÁCH API CHÍNH",
        "PHỤ LỤC C. DỮ LIỆU DEMO TỪ SEED SCRIPT",
        "PHỤ LỤC D. NHẬN XÉT VỀ MỨC ĐỘ HOÀN THIỆN HIỆN TẠI",
    ]
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)
    doc.add_page_break()


def add_markdown_table(doc, rows):
    parsed = []
    for row in rows:
        cells = [cell.strip() for cell in row.strip().strip("|").split("|")]
        parsed.append(cells)
    if len(parsed) >= 2 and all(re.fullmatch(r":?-{3,}:?", c.replace(" ", "")) for c in parsed[1]):
        parsed.pop(1)
    if not parsed:
        return
    max_cols = max(len(row) for row in parsed)
    widths = [PAGE_WIDTH_DXA // max_cols for _ in range(max_cols)]
    widths[-1] += PAGE_WIDTH_DXA - sum(widths)
    table = doc.add_table(rows=len(parsed), cols=max_cols)
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    for r_idx, row in enumerate(parsed):
        for c_idx in range(max_cols):
            cell = table.rows[r_idx].cells[c_idx]
            text = row[c_idx] if c_idx < len(row) else ""
            cell.text = ""
            p = cell.paragraphs[0]
            set_paragraph_spacing(p, after=2, line=1.08)
            run = p.add_run(text)
            set_run_font(run, size=8.5 if max_cols >= 4 else 9.5, bold=(r_idx == 0))
            if r_idx == 0:
                set_cell_shading(cell, "F2F4F7")
    spacer = doc.add_paragraph()
    set_paragraph_spacing(spacer, before=0, after=4)


def add_markdown(doc, markdown):
    lines = markdown.strip().splitlines()
    table_lines = []
    def flush_table():
        nonlocal table_lines
        if table_lines:
            add_markdown_table(doc, table_lines)
            table_lines = []
    for line in lines:
        line = line.rstrip()
        if line.strip().startswith("|") and line.strip().endswith("|"):
            table_lines.append(line)
            continue
        flush_table()
        if not line.strip():
            continue
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith("#### "):
            p = doc.add_paragraph()
            set_paragraph_spacing(p, before=6, after=3)
            run = p.add_run(line[5:].strip())
            set_run_font(run, size=11.5, color=RGBColor(31, 77, 120), bold=True)
        elif line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(line[2:].strip())
        elif re.match(r"^\d+\.\s+", line):
            p = doc.add_paragraph(style="List Number")
            p.add_run(re.sub(r"^\d+\.\s+", "", line).strip())
        elif line.startswith("@") or line.startswith("actor ") or line.startswith("usecase ") or line.startswith("rectangle ") or " --> " in line or line == "}":
            p = doc.add_paragraph()
            set_paragraph_spacing(p, after=0, line=1.0)
            run = p.add_run(line)
            set_run_font(run, name="Courier New", size=8.5, color=RGBColor(45, 45, 45))
        else:
            p = doc.add_paragraph()
            set_paragraph_spacing(p, after=6, line=1.10)
            run = p.add_run(line.strip())
            set_run_font(run, size=11)
    flush_table()


def main():
    OUT_DIR.mkdir(exist_ok=True)
    markdown = SOURCE_MD.read_text(encoding="utf-8")
    txt = REPORT_TITLE + "\n" + REPORT_SUBTITLE + "\n\n"
    for label, value in REPORT_META:
        txt += f"{label}: {value}\n"
    txt += "\n" + markdown.strip() + "\n"
    TXT_PATH.write_text(txt, encoding="utf-8")
    doc = Document()
    configure_document(doc)
    add_cover(doc)
    add_static_toc(doc)
    add_markdown(doc, markdown)
    doc.save(DOCX_PATH)
    print(DOCX_PATH)
    print(TXT_PATH)


if __name__ == "__main__":
    main()
